"""
Document Parser Service
Handles parsing of various document formats (PDF, Word) to extract text content
Provides comprehensive logging and error handling
"""

import logging
import io
from typing import Optional, Dict, Any
from pathlib import Path

# Import file parsing libraries
try:
    import PyPDF2
    import pdfplumber
    from docx import Document as DocxDocument
    PARSING_AVAILABLE = True
except ImportError as e:
    PARSING_AVAILABLE = False
    IMPORT_ERROR = str(e)

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Service for parsing document files and extracting text content
    Supports PDF and Word (.docx) formats
    """
    
    def __init__(self):
        """Initialize the document parser"""
        if not PARSING_AVAILABLE:
            logger.error(
                '[DocumentParser] File parsing libraries not available',
                extra={
                    'error': IMPORT_ERROR,
                    'required_packages': ['PyPDF2', 'pdfplumber', 'python-docx']
                }
            )
        else:
            logger.info('[DocumentParser] Document parser initialized successfully')
    
    def is_available(self) -> bool:
        """Check if parsing libraries are available"""
        return PARSING_AVAILABLE
    
    def parse_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse a file and extract its text content
        
        Args:
            file_content: Binary content of the file
            filename: Name of the file (used to determine file type)
        
        Returns:
            Dictionary with parsing result:
            {
                'success': bool,
                'text': str (extracted text content),
                'error': str (error message if failed),
                'metadata': dict (file metadata)
            }
        """
        logger.info(
            '[DocumentParser] Starting file parsing',
            extra={
                'file_name': filename,
                'file_size': len(file_content),
                'file_type': Path(filename).suffix.lower()
            }
        )
        
        if not PARSING_AVAILABLE:
            error_msg = 'File parsing libraries not available'
            logger.error(
                f'[DocumentParser] {error_msg}',
                extra={'file_name': filename}
            )
            return {
                'success': False,
                'text': '',
                'error': error_msg,
                'metadata': {}
            }
        
        try:
            # Determine file type from extension
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.pdf':
                return self._parse_pdf(file_content, filename)
            elif file_ext in ['.docx', '.doc']:
                return self._parse_word(file_content, filename)
            else:
                error_msg = f'Unsupported file type: {file_ext}'
                logger.warning(
                    f'[DocumentParser] {error_msg}',
                    extra={
                        'file_name': filename,
                        'file_type': file_ext,
                        'supported_types': ['.pdf', '.docx', '.doc']
                    }
                )
                return {
                    'success': False,
                    'text': '',
                    'error': error_msg,
                    'metadata': {'filename': filename}
                }
        
        except Exception as e:
            error_msg = f'Error parsing file: {str(e)}'
            logger.error(
                f'[DocumentParser] {error_msg}',
                extra={
                    'filename': filename,
                    'error_type': type(e).__name__,
                    'error_message': str(e)
                },
                exc_info=True
            )
            return {
                'success': False,
                'text': '',
                'error': error_msg,
                'metadata': {'filename': filename}
            }
    
    def _parse_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse PDF file and extract text content
        Uses pdfplumber as primary parser, falls back to PyPDF2 if needed
        
        Args:
            file_content: Binary content of the PDF file
            filename: Name of the file
        
        Returns:
            Dictionary with parsing result
        """
        logger.info(
            '[DocumentParser] Parsing PDF file',
            extra={'file_name': filename}
        )
        
        extracted_text = ''
        page_count = 0
        parsing_method = 'unknown'
        
        try:
            # Method 1: Try pdfplumber first (better text extraction)
            try:
                logger.debug(
                    '[DocumentParser] Attempting PDF parsing with pdfplumber',
                    extra={'file_name': filename}
                )
                
                file_obj = io.BytesIO(file_content)
                with pdfplumber.open(file_obj) as pdf:
                    page_count = len(pdf.pages)
                    logger.info(
                        '[DocumentParser] PDF opened successfully with pdfplumber',
                        extra={
                            'file_name': filename,
                            'page_count': page_count
                        }
                    )
                    
                    for i, page in enumerate(pdf.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                extracted_text += page_text + '\n\n'
                                logger.debug(
                                    '[DocumentParser] Extracted text from page',
                                    extra={
                                        'filename': filename,
                                        'page_number': i,
                                        'text_length': len(page_text)
                                    }
                                )
                        except Exception as page_error:
                            logger.warning(
                                '[DocumentParser] Failed to extract text from page',
                                extra={
                                    'file_name': filename,
                                    'page_number': i,
                                    'error': str(page_error)
                                }
                            )
                
                parsing_method = 'pdfplumber'
                logger.info(
                    '[DocumentParser] PDF parsed successfully with pdfplumber',
                    extra={
                        'file_name': filename,
                        'page_count': page_count,
                        'total_text_length': len(extracted_text),
                        'method': parsing_method
                    }
                )
            
            except Exception as pdfplumber_error:
                # Method 2: Fall back to PyPDF2
                logger.warning(
                    '[DocumentParser] pdfplumber parsing failed, falling back to PyPDF2',
                    extra={
                        'file_name': filename,
                        'error': str(pdfplumber_error)
                    }
                )
                
                file_obj = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(file_obj)
                page_count = len(pdf_reader.pages)
                
                logger.info(
                    '[DocumentParser] PDF opened successfully with PyPDF2',
                    extra={
                        'file_name': filename,
                        'page_count': page_count
                    }
                )
                
                for i, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            extracted_text += page_text + '\n\n'
                            logger.debug(
                                '[DocumentParser] Extracted text from page',
                                extra={
                                    'file_name': filename,
                                    'page_number': i,
                                    'text_length': len(page_text)
                                }
                            )
                    except Exception as page_error:
                        logger.warning(
                            '[DocumentParser] Failed to extract text from page',
                            extra={
                                'filename': filename,
                                'page_number': i,
                                'error': str(page_error)
                            }
                        )
                
                parsing_method = 'PyPDF2'
                logger.info(
                    '[DocumentParser] PDF parsed successfully with PyPDF2',
                    extra={
                        'file_name': filename,
                        'page_count': page_count,
                        'total_text_length': len(extracted_text),
                        'method': parsing_method
                    }
                )
            
            # Clean up extracted text
            extracted_text = extracted_text.strip()
            
            if not extracted_text:
                logger.warning(
                    '[DocumentParser] No text extracted from PDF',
                    extra={
                        'file_name': filename,
                        'page_count': page_count,
                        'method': parsing_method
                    }
                )
                return {
                    'success': False,
                    'text': '',
                    'error': 'No text content found in PDF file',
                    'metadata': {
                        'filename': filename,
                        'page_count': page_count,
                        'parsing_method': parsing_method
                    }
                }
            
            logger.info(
                '[DocumentParser] PDF parsing completed successfully',
                extra={
                    'file_name': filename,
                    'page_count': page_count,
                    'text_length': len(extracted_text),
                    'parsing_method': parsing_method
                }
            )
            
            return {
                'success': True,
                'text': extracted_text,
                'error': None,
                'metadata': {
                    'filename': filename,
                    'page_count': page_count,
                    'text_length': len(extracted_text),
                    'parsing_method': parsing_method
                }
            }
        
        except Exception as e:
            error_msg = f'Error parsing PDF: {str(e)}'
            logger.error(
                f'[DocumentParser] {error_msg}',
                extra={
                    'file_name': filename,
                    'error_type': type(e).__name__,
                    'error_message': str(e)
                },
                exc_info=True
            )
            return {
                'success': False,
                'text': '',
                'error': error_msg,
                'metadata': {'filename': filename}
            }
    
    def _parse_word(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse Word (.docx) file and extract text content
        
        Args:
            file_content: Binary content of the Word file
            filename: Name of the file
        
        Returns:
            Dictionary with parsing result
        """
        logger.info(
            '[DocumentParser] Parsing Word file',
            extra={'file_name': filename}
        )
        
        try:
            file_obj = io.BytesIO(file_content)
            document = DocxDocument(file_obj)
            
            # Count paragraphs
            paragraph_count = len(document.paragraphs)
            logger.info(
                '[DocumentParser] Word document opened successfully',
                extra={
                    'file_name': filename,
                    'paragraph_count': paragraph_count
                }
            )
            
            # Extract text from all paragraphs
            extracted_text = ''
            non_empty_paragraphs = 0
            
            for i, paragraph in enumerate(document.paragraphs, 1):
                text = paragraph.text.strip()
                if text:
                    extracted_text += text + '\n\n'
                    non_empty_paragraphs += 1
                    logger.debug(
                        '[DocumentParser] Extracted text from paragraph',
                        extra={
                            'file_name': filename,
                            'paragraph_number': i,
                            'text_length': len(text)
                        }
                    )
            
            # Clean up extracted text
            extracted_text = extracted_text.strip()
            
            if not extracted_text:
                logger.warning(
                    '[DocumentParser] No text extracted from Word document',
                    extra={
                        'file_name': filename,
                        'paragraph_count': paragraph_count,
                        'non_empty_paragraphs': non_empty_paragraphs
                    }
                )
                return {
                    'success': False,
                    'text': '',
                    'error': 'No text content found in Word document',
                    'metadata': {
                        'filename': filename,
                        'paragraph_count': paragraph_count,
                        'non_empty_paragraphs': non_empty_paragraphs
                    }
                }
            
            logger.info(
                '[DocumentParser] Word document parsing completed successfully',
                extra={
                    'file_name': filename,
                    'paragraph_count': paragraph_count,
                    'non_empty_paragraphs': non_empty_paragraphs,
                    'text_length': len(extracted_text)
                }
            )
            
            return {
                'success': True,
                'text': extracted_text,
                'error': None,
                'metadata': {
                    'filename': filename,
                    'paragraph_count': paragraph_count,
                    'non_empty_paragraphs': non_empty_paragraphs,
                    'text_length': len(extracted_text)
                }
            }
        
        except Exception as e:
            error_msg = f'Error parsing Word document: {str(e)}'
            logger.error(
                f'[DocumentParser] {error_msg}',
                extra={
                    'file_name': filename,
                    'error_type': type(e).__name__,
                    'error_message': str(e)
                },
                exc_info=True
            )
            return {
                'success': False,
                'text': '',
                'error': error_msg,
                'metadata': {'filename': filename}
            }


# Create global parser instance
document_parser = DocumentParser()

